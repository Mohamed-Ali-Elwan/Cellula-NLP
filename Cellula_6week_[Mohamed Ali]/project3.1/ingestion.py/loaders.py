from pathlib import Path

from langchain_core.documents import Document


class FileLoader:

    @staticmethod
    def load(path: str) -> list[Document]:

        file_path = Path(path)
        extension = file_path.suffix.lower()

        if extension == ".pdf":
            return FileLoader._load_pdf(file_path)

        elif extension == ".docx":
            return FileLoader._load_docx(file_path)

        elif extension in {
            ".txt",
            ".md",
            ".py",
            ".java",
            ".cpp",
            ".c",
            ".h",
            ".js",
            ".ts",
        }:
            return FileLoader._load_text(file_path)


        elif extension == ".wav":
            return FileLoader._load_audio(file_path)

        else:
            raise ValueError(
                f"Unsupported file type: {extension}"
            )

    # --------------------------------------------------
    # PDF
    # --------------------------------------------------

    @staticmethod
    def _load_pdf(path: Path) -> list[Document]:

        from pypdf import PdfReader

        reader = PdfReader(str(path))

        documents = []

        for page_number, page in enumerate(
            reader.pages,
            start=1
        ):

            text = page.extract_text() or ""

            if text.strip():

                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": str(path),
                            "type": "pdf",
                            "page": page_number,
                        },
                    )
                )

        return documents

    # --------------------------------------------------
    # DOCX
    # --------------------------------------------------

    @staticmethod
    def _load_docx(path: Path) -> list[Document]:

        from docx import Document as DocxDocument

        document = DocxDocument(str(path))

        text = "\n".join(
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        )

        return [
            Document(
                page_content=text,
                metadata={
                    "source": str(path),
                    "type": "docx",
                },
            )
        ]

    # --------------------------------------------------
    # TXT / Markdown / Code
    # --------------------------------------------------

    @staticmethod
    def _load_text(path: Path) -> list[Document]:

        text = path.read_text(
            encoding="utf-8",
            errors="ignore",
        )

        return [
            Document(
                page_content=text,
                metadata={
                    "source": str(path),
                    "type": path.suffix.lower(),
                },
            )
        ]

    

    # --------------------------------------------------
    # WAV
    # --------------------------------------------------

    @staticmethod
    def _load_audio(path: Path) -> list[Document]:

        from faster_whisper import WhisperModel

        model = WhisperModel(
            "base",
            compute_type="int8",
        )

        segments, _ = model.transcribe(
            str(path)
        )

        text = " ".join(
            segment.text
            for segment in segments
        )

        return [
            Document(
                page_content=text,
                metadata={
                    "source": str(path),
                    "type": "wav",
                },
            )
        ]